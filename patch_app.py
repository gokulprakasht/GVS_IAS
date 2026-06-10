import sys
from pathlib import Path

app_path = Path(r"C:\IAS\IAS_CLOUD_FULL\app.py")
if not app_path.exists():
    print("ERROR: app.py not found"); sys.exit(1)

content = app_path.read_text(encoding="utf-8")
fixes = 0

# FIX 1: API key direct inject (bypass read-only filesystem on Streamlit Cloud)
old1 = '''        # API key
        _ak = _sec.get("ANTHROPIC_API_KEY","")
        if _ak and _ak.startswith("sk-ant-"):
            _kf = ROOT / "api_key.txt"
            if not _kf.exists() or _kf.read_text().strip() != _ak:
                _kf.write_text(_ak)'''
new1 = '''        # API key — inject directly into memory (read-only filesystem safe)
        _ak = _sec.get("ANTHROPIC_API_KEY","")
        if _ak and _ak.startswith("sk-ant-"):
            apikey._KEY = _ak
            try:
                _kf = ROOT / "api_key.txt"
                _kf.write_text(_ak)
            except Exception:
                pass'''
if old1 in content:
    content = content.replace(old1, new1, 1)
    fixes += 1
    print("FIX 1 applied: API key injected directly into memory")
else:
    if 'apikey._KEY = _ak' in content:
        print("FIX 1 already applied")
    else:
        print("FIX 1 FAILED: pattern not found")

# FIX 2: _extract_text use BytesIO
old2 = '''def _extract_text(f):
    import tempfile as tf, os as _os
    nm = f.name.lower()
    # Seek to start — Streamlit may have read the buffer already
    try: f.seek(0)
    except Exception: pass
    raw = f.read()
    if not raw:
        try: f.seek(0); raw = f.read()
        except: pass
    with tf.NamedTemporaryFile(delete=False, suffix=_os.path.splitext(nm)[1]) as t:
        t.write(raw); tp = t.name
    try:
        if nm.endswith(".pdf"):
            from pypdf import PdfReader
            return " ".join(p.extract_text() or "" for p in PdfReader(tp).pages).strip()
        elif nm.endswith(".docx"):
            from docx import Document
            return "\\n".join(p.text for p in Document(tp).paragraphs if p.text.strip())
        elif nm.endswith(".txt"):
            return raw.decode("utf-8", "replace").strip()
    except Exception as e: return f"Error: {e}"
    finally:
        try: _os.unlink(tp)
        except: pass
    return ""'''
new2 = '''def _extract_text(f):
    from io import BytesIO
    nm = f.name.lower()
    try: f.seek(0)
    except Exception: pass
    raw = f.read()
    if not raw:
        try: f.seek(0); raw = f.read()
        except: pass
    if not raw:
        return "Error: file is empty"
    bio = BytesIO(raw)
    try:
        if nm.endswith(".pdf"):
            try:
                from pypdf import PdfReader
            except ImportError:
                from PyPDF2 import PdfReader
            return " ".join(p.extract_text() or "" for p in PdfReader(bio).pages).strip()
        elif nm.endswith(".docx"):
            from docx import Document
            return "\\n".join(p.text for p in Document(bio).paragraphs if p.text.strip())
        elif nm.endswith(".txt"):
            return raw.decode("utf-8", "replace").strip()
        else:
            return "Error: unsupported file type"
    except Exception as e:
        return f"Error: {e}"'''
if old2 in content:
    content = content.replace(old2, new2, 1)
    fixes += 1
    print("FIX 2 applied: CV upload uses BytesIO")
else:
    if 'from io import BytesIO' in content and 'def _extract_text' in content:
        print("FIX 2 already applied")
    else:
        print("FIX 2 FAILED: pattern not found")

# FIX 3: settings_e defined before report preview
old3 = '                st.markdown("#### \U0001f4c4 Generate DOCX Report")\n                # \u2500\u2500 REPORT PREVIEW \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\n                with st.expander("\U0001f441 Report Preview", expanded=False):'
new3 = '                st.markdown("#### \U0001f4c4 Generate DOCX Report")\n                settings_e = cfg.get_settings()\n                # \u2500\u2500 REPORT PREVIEW \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\n                with st.expander("\U0001f441 Report Preview", expanded=False):'
if old3 in content:
    content = content.replace(old3, new3, 1)
    fixes += 1
    print("FIX 3 applied: settings_e defined before use")
else:
    if 'settings_e = cfg.get_settings()' in content:
        print("FIX 3 already applied")
    else:
        print("FIX 3 FAILED: pattern not found")

app_path.write_text(content, encoding="utf-8")
print(f"\nTotal fixes applied: {fixes}")
print("app.py patched successfully")
